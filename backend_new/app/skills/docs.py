"""Document skills: create/edit Word docs and text files.

Word docs use python-docx when installed, with a plain-text fallback so
the skill always works.
"""

import os
from typing import Any, Dict


def _try_import_docx():
    try:
        import docx
        return docx
    except Exception:
        return None


async def docs_new(params: Dict[str, Any], ctx: Any) -> Dict[str, Any]:
    kind = (params.get("kind") or "word").lower()
    name = params.get("name") or "jarvis_document"
    data_dir = getattr(ctx, "kernel", None) and ctx.kernel.cfg.data_dir or os.getcwd()
    os.makedirs(data_dir, exist_ok=True)

    if kind in ("word", "doc", "docx"):
        path = os.path.join(str(data_dir), f"{name}.docx")
        docx = _try_import_docx()
        if docx is not None:
            doc = docx.Document()
            doc.add_heading(name.replace("_", " ").title(), 0)
            doc.add_paragraph("Created by J.A.R.V.I.S.")
            doc.save(path)
        else:
            path = path[:-5] + ".txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(f"{name}\n\nCreated by J.A.R.V.I.S. (python-docx not installed)\n")
        return {"success": True, "narration": f"Created document {os.path.basename(path)}.",
                "type": "docs_result", "data": {"path": path, "kind": kind}}

    path = os.path.join(str(data_dir), f"{name}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write("New document.\n")
    return {"success": True, "narration": f"Created text file {os.path.basename(path)}.",
            "type": "docs_result", "data": {"path": path, "kind": "text"}}


async def docs_edit(params: Dict[str, Any], ctx: Any) -> Dict[str, Any]:
    data_dir = getattr(ctx, "kernel", None) and ctx.kernel.cfg.data_dir or os.getcwd()
    path = params.get("path") or os.path.join(str(data_dir), "jarvis_document.txt")
    if not os.path.exists(path):
        return {"success": False, "narration": f"Document '{os.path.basename(path)}' not found.",
                "type": "docs_result"}
    if os.name == "nt":
        os.startfile(path)
        return {"success": True, "narration": f"Opening {os.path.basename(path)} for editing.",
                "type": "docs_result", "data": {"path": path}}
    return {"success": False, "narration": "Auto-edit needs a text of changes to apply.",
            "type": "docs_result"}


def register(reg) -> None:
    reg.skill("docs_new", docs_new, description="Create a new Word document")
    reg.skill("docs_edit", docs_edit, description="Open a document for editing")
